
'''
import docx

docx_patch = ".\data\mot_202122__4040.docx"
document = docx.Document(docx_patch)
tables = document.tables
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text)

'''

import tabula

pdf_path = ".\data\mot_202122__4040.pdf"

dfs = tabula.read_pdf(pdf_path, pages="20", lattice = True)

print (len(dfs))

print (dfs[0])